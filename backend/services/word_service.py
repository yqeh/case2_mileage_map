"""Word report generation service."""

from datetime import datetime
from pathlib import Path
import os
from urllib.parse import urlencode

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from loguru import logger
from PIL import Image, ImageDraw, ImageFont

from services.gmap_screenshot_service import capture_maps_url_screenshot_sync
from services.google_maps_service import GoogleMapsService
from utils.log_sanitizer import sanitize_filename
from utils.path_manager import get_base_dir, get_output_dir, get_relative_path, get_temp_maps_dir


class WordService:
    """Generate Word reports for travel records."""

    def __init__(self):
        self.output_dir = get_output_dir()
        self.maps_service = GoogleMapsService()

    def _format_mmdd(self, date_value):
        try:
            if not date_value:
                return ""
            if isinstance(date_value, datetime):
                dt = date_value
            elif isinstance(date_value, str):
                try:
                    dt = datetime.strptime(date_value, '%Y-%m-%d')
                except ValueError:
                    dt = datetime.fromisoformat(date_value)
            else:
                return str(date_value)
            return f"{dt.month}/{dt.day}"
        except Exception:
            return str(date_value)

    def _safe_dt(self, date_value):
        if not date_value:
            return datetime.min
        if isinstance(date_value, datetime):
            return date_value
        if isinstance(date_value, str):
            try:
                return datetime.fromisoformat(date_value)
            except ValueError:
                try:
                    return datetime.strptime(date_value, "%Y-%m-%d")
                except ValueError:
                    return datetime.min
        return datetime.min

    def _load_font(self, size: int):
        candidates = [
            get_base_dir() / 'assets' / 'fonts' / 'NotoSansTC-Regular.ttf',
            Path(os.environ.get('WINDIR', 'C:/Windows')) / 'Fonts' / 'msjh.ttc',
            Path(os.environ.get('WINDIR', 'C:/Windows')) / 'Fonts' / 'mingliu.ttc',
        ]
        for fp in candidates:
            try:
                if str(fp).lower().endswith('.ttc'):
                    return ImageFont.truetype(str(fp), size, index=0)
                return ImageFont.truetype(str(fp), size)
            except Exception:
                continue
        return ImageFont.load_default()

    def _resolve_existing_image(self, map_image_path):
        if not map_image_path:
            return None
        base_dir = get_base_dir()
        clean_path = str(map_image_path).lstrip('/\\')
        absolute_image_path = base_dir / clean_path
        if not absolute_image_path.exists():
            return None
        if os.path.getsize(absolute_image_path) <= 10240:
            return None
        return absolute_image_path

    def _stamp_timestamp(self, image_path: Path):
        now_text = datetime.now().strftime('截圖時間: %Y/%m/%d %H:%M')
        image = Image.open(image_path).convert('RGB')
        draw = ImageDraw.Draw(image)
        font = self._load_font(44)
        bbox = draw.textbbox((0, 0), now_text, font=font)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]
        pad = 20
        x = image.width - text_w - pad * 2 - 28
        y = image.height - text_h - pad * 2 - 28
        draw.rounded_rectangle((x, y, x + text_w + pad * 2, y + text_h + pad * 2), radius=18, fill=(255, 255, 255))
        draw.text((x + pad, y + pad), now_text, fill=(40, 40, 40), font=font)
        image.save(image_path)

    def _capture_image_from_link(self, record):
        maps_url = record.get('連結') or record.get('GoogleMapUrl') or record.get('google_map_url')
        if not maps_url:
            return None
        temp_maps_dir = get_temp_maps_dir()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]
        output_path = temp_maps_dir / f'word_link_{timestamp}.png'
        screenshot_path = capture_maps_url_screenshot_sync(
            maps_url=maps_url,
            output_path=str(output_path),
            viewport_width=1366,
            viewport_height=768,
            wait_timeout=30000,
            log_context=record.get('目的地名稱') or record.get('終點地址') or 'Google Maps link',
        )
        if not screenshot_path:
            return None
        absolute_path = Path(screenshot_path)
        if not absolute_path.exists() or os.path.getsize(absolute_path) <= 10240:
            return None
        self._stamp_timestamp(absolute_path)
        relative_path = get_relative_path(absolute_path)
        if not relative_path.startswith('/'):
            relative_path = '/' + relative_path
        record['StaticMapImage'] = relative_path
        return absolute_path

    def _pick_text(self, *values):
        for value in values:
            text = str(value or '').strip()
            if text:
                return text
        return ''

    def _capture_image_from_route(self, origin_address, destination_address, record):
        if not origin_address or not destination_address:
            return None
        maps_url = self._build_multi_stop_maps_url(origin_address, [destination_address], origin_address)
        return self._capture_image_from_maps_url(
            maps_url,
            record,
            log_context=f'{origin_address} -> {destination_address} -> {origin_address}',
        )

    def _capture_image_from_maps_url(self, maps_url, record, log_context=None, metadata=None):
        if not maps_url:
            return None
        temp_maps_dir = get_temp_maps_dir()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]
        output_path = temp_maps_dir / f'word_route_{timestamp}.png'
        screenshot_path = capture_maps_url_screenshot_sync(
            maps_url=maps_url,
            output_path=str(output_path),
            viewport_width=1366,
            viewport_height=768,
            wait_timeout=30000,
            log_context=log_context or maps_url,
            metadata=metadata,
        )
        if not screenshot_path:
            return None
        absolute_path = Path(screenshot_path)
        if not absolute_path.exists() or os.path.getsize(absolute_path) <= 10240:
            return None
        self._stamp_timestamp(absolute_path)
        relative_path = get_relative_path(absolute_path)
        if not relative_path.startswith('/'):
            relative_path = '/' + relative_path
        if record is not None:
            record['StaticMapImage'] = relative_path
        return absolute_path

    def _build_multi_stop_maps_url(self, origin, stops, destination):
        clean_stops = [str(stop).strip() for stop in stops if str(stop or '').strip()]
        query = {
            'api': '1',
            'origin': str(origin or '').strip(),
            'destination': str(destination or '').strip(),
            'travelmode': 'driving',
        }
        if clean_stops:
            query['waypoints'] = '|'.join(clean_stops)
        return 'https://www.google.com/maps/dir/?' + urlencode(query)

    def _format_km(self, value):
        if value is None:
            return '--'
        try:
            number = float(value)
        except (TypeError, ValueError):
            return '--'
        if number < 0:
            return '--'
        return str(int(number))

    def _calculate_route_total_km(self, route_points, records=None):
        total = 0.0
        has_distance = False
        for origin, destination in zip(route_points, route_points[1:]):
            if not origin or not destination:
                continue
            detail = self.maps_service.get_route_detail(origin, destination, alternatives=False)
            if not detail.get('success'):
                logger.warning(f"\u6838\u92b7\u516c\u91cc\u6578\u8a08\u7b97\u5931\u6557: {origin} -> {destination}, {detail.get('error')}")
                continue
            total += float(detail.get('distance_km') or 0)
            has_distance = True
        if has_distance:
            return round(total, 2)
        return self._estimate_total_km_from_records(records or [])

    def _estimate_total_km_from_records(self, records):
        total = 0.0
        has_distance = False
        for record in records:
            value = record.get('OneWayKm')
            if value is None:
                round_trip = record.get('RoundTripKm')
                try:
                    value = float(round_trip) / 2 if round_trip is not None else None
                except (TypeError, ValueError):
                    value = None
            try:
                if value is not None:
                    total += float(value)
                    has_distance = True
            except (TypeError, ValueError):
                continue
        if not has_distance:
            return None
        return round(total * 2, 2)

    def _date_key(self, date_value):
        dt = self._safe_dt(date_value)
        if dt == datetime.min:
            return str(date_value or '')
        return dt.date().isoformat()

    def _trip_group_key(self, record):
        return (
            self._date_key(record.get('\u51fa\u5dee\u65e5\u671f\u6642\u9593\uff08\u958b\u59cb\uff09')),
            str(record.get('\u90e8\u9580') or '').strip(),
            str(record.get('\u59d3\u540d') or '').strip(),
        )

    def _group_records_by_day(self, records):
        sorted_records = sorted(records, key=lambda x: self._safe_dt(x.get('\u51fa\u5dee\u65e5\u671f\u6642\u9593\uff08\u958b\u59cb\uff09')))
        groups = []
        group_index = {}
        for record in sorted_records:
            key = self._trip_group_key(record)
            if key not in group_index:
                group_index[key] = len(groups)
                groups.append([])
            groups[group_index[key]].append(record)
        return groups

    def _resolve_trip_plan(self, records, fixed_origin):
        first_record = records[0]
        start_name = self._pick_text(first_record.get('\u8d77\u9ede\u540d\u7a31'))
        start_address = self._pick_text(
            first_record.get('\u8d77\u9ede\u5730\u5740'),
            first_record.get('OriginAddress'),
            first_record.get('origin_address'),
        )
        company_display = self._pick_text(start_name, fixed_origin, start_address)
        company_route = self._pick_text(fixed_origin, start_address, start_name)

        stops = []
        for record in records:
            destination_name = self._pick_text(record.get('\u76ee\u7684\u5730\u540d\u7a31'))
            destination_address = self._pick_text(
                record.get('\u7d42\u9ede\u5730\u5740'),
                record.get('DestinationAddress'),
                record.get('destination_address'),
            )
            destination_display = self._pick_text(destination_name, destination_address)
            destination_route = self._pick_text(destination_address, destination_name)
            if destination_route:
                stops.append({
                    'display': destination_display,
                    'route': destination_route,
                    'record': record,
                })

        return company_display, company_route, stops

    def generate_report(self, project_name, records, fixed_origin=None, page_break_per_record=True):
        try:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(0.25)
                section.bottom_margin = Inches(0.25)
                section.left_margin = Inches(0.25)
                section.right_margin = Inches(0.25)
            record_groups = self._group_records_by_day(records)

            for idx, group_records in enumerate(record_groups):
                try:
                    logger.info(f"\u8655\u7406\u7b2c {idx + 1}/{len(record_groups)} \u7d44\u540c\u65e5\u884c\u7a0b")
                    if idx > 0 and page_break_per_record and idx % 2 == 0:
                        doc.add_page_break()

                    date_str = self._format_mmdd(group_records[0].get('\u51fa\u5dee\u65e5\u671f\u6642\u9593\uff08\u958b\u59cb\uff09'))
                    company_display, company_route, stops = self._resolve_trip_plan(group_records, fixed_origin)
                    if not stops:
                        logger.warning(f"\u7b2c {idx + 1} \u7d44\u6c92\u6709\u76ee\u7684\u5730\uff0c\u7565\u904e")
                        continue

                    destination_displays = [stop['display'] for stop in stops]
                    destination_routes = [stop['route'] for stop in stops]
                    route_points = [company_route, *destination_routes, company_route]
                    total_km = self._calculate_route_total_km(route_points, group_records)
                    final_destination_display = destination_displays[-1]

                    maps_url = self._build_multi_stop_maps_url(company_route, destination_routes, company_route)
                    first_record = group_records[0]
                    capture_metadata = {}
                    absolute_image_path = self._capture_image_from_maps_url(
                        maps_url,
                        first_record,
                        log_context=f"{company_route} -> {' -> '.join(destination_routes)} -> {company_route}",
                        metadata=capture_metadata,
                    )
                    if total_km is None:
                        total_km = capture_metadata.get('distance_km')
                    if not absolute_image_path and len(group_records) == 1:
                        absolute_image_path = self._capture_image_from_link(first_record)

                    total_km_text = self._format_km(total_km)
                    if len(destination_displays) == 1:
                        title_text = f"{date_str}{company_display}\u81f3{final_destination_display}\u5f80\u8fd4\uff0c \u5171\u6838\u92b7 {total_km_text} \u516c\u91cc"
                    else:
                        joined_destinations = '\u3001'.join(destination_displays)
                        title_text = f"{date_str}{company_display}\u81f3{joined_destinations}\u5f80\u8fd4\uff0c \u5171\u6838\u92b7 {total_km_text} \u516c\u91cc"

                    title_paragraph = doc.add_paragraph(title_text)
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    title_paragraph.paragraph_format.space_before = Pt(0)
                    title_paragraph.paragraph_format.space_after = Pt(0)
                    title_paragraph.paragraph_format.line_spacing = 1
                    for run in title_paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(18)

                    if absolute_image_path:
                        picture_paragraph = doc.add_paragraph()
                        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        picture_paragraph.paragraph_format.space_before = Pt(0)
                        picture_paragraph.paragraph_format.space_after = Pt(0)
                        picture_paragraph.paragraph_format.line_spacing = 1
                        run = picture_paragraph.add_run()
                        run.add_picture(str(absolute_image_path), width=Inches(7.95))
                    else:
                        error_paragraph = doc.add_paragraph()
                        error_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        error_run = error_paragraph.add_run('\u672c\u7d44\u5730\u5716\u622a\u5716\u5931\u6557')
                        error_run.font.size = Pt(14)
                except Exception as e:
                    logger.error(f"\u8655\u7406\u7b2c {idx + 1} \u7d44\u540c\u65e5\u884c\u7a0b\u6642\u767c\u751f\u932f\u8aa4: {e}")
                    continue

            project_display_name = project_name or '未分類'
            filename = sanitize_filename(f"{project_display_name}_里程報表.docx") or '未分類_里程報表.docx'
            file_path = self.output_dir / filename
            doc.save(str(file_path))
            logger.info(f"報表已儲存: {str(file_path)}")
            return str(file_path)
        except Exception as e:
            logger.error(f"產生 Word 報表錯誤: {str(e)}")
            raise
